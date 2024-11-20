from fastapi import FastAPI, UploadFile, File, HTTPException, Depends
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import List, Optional, Dict
import uuid
import os
import tempfile
from datetime import datetime

from langchain_core.messages import HumanMessage
from langchain_core.documents import Document
from langchain_core.output_parsers import StrOutputParser
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.runnables import RunnableLambda, RunnablePassthrough

from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain.retrievers.multi_vector import MultiVectorRetriever
from langchain.storage import InMemoryStore
from langchain_chroma import Chroma
from langchain_text_splitters import CharacterTextSplitter
from unstructured.partition.pdf import partition_pdf
from dotenv import load_dotenv

app = FastAPI(title="Multimodal RAG API")

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Pydantic models
class ChatMessage(BaseModel):
    role: str
    content: str
    timestamp: datetime

class Conversation(BaseModel):
    conversation_id: str
    messages: List[ChatMessage]
    
class QueryRequest(BaseModel):
    conversation_id: str
    query: str

class QueryResponse(BaseModel):
    response: str
    retrieved_documents: List[str]

# In-memory storage
conversations: Dict[str, Conversation] = {}
document_processors = {
    "pdf": extract_pdf_elements,
    "csv": process_csv,  # You'll need to implement this
    "txt": process_text,  # You'll need to implement this
    "doc": process_doc,   # You'll need to implement this
    "image": process_image # You'll need to implement this
}

retrievers: Dict[str, MultiVectorRetriever] = {}

# Dependency for getting conversation
async def get_conversation(conversation_id: str) -> Conversation:
    if conversation_id not in conversations:
        raise HTTPException(status_code=404, detail="Conversation not found")
    return conversations[conversation_id]

@app.post("/conversations", response_model=Conversation)
async def create_conversation():
    """Create a new conversation"""
    conversation_id = str(uuid.uuid4())
    conversation = Conversation(
        conversation_id=conversation_id,
        messages=[]
    )
    conversations[conversation_id] = conversation
    return conversation

@app.post("/upload/{conversation_id}")
async def upload_files(
    conversation_id: str,
    files: List[UploadFile] = File(...),
    conversation: Conversation = Depends(get_conversation)
):
    """Upload and process multiple files"""
    try:
        with tempfile.TemporaryDirectory() as temp_dir:
            # Process each file based on its type
            for file in files:
                file_extension = file.filename.split('.')[-1].lower()
                temp_file_path = os.path.join(temp_dir, file.filename)
                
                # Save uploaded file
                with open(temp_file_path, "wb") as f:
                    content = await file.read()
                    f.write(content)
                
                # Process file based on type
                if file_extension in document_processors:
                    processor = document_processors[file_extension]
                    elements = processor(temp_file_path, temp_dir)
                    
                    # Process elements similar to original code
                    texts, tables = categorize_elements(elements)
                    text_splitter = CharacterTextSplitter.from_tiktoken_encoder(
                        chunk_size=4000, chunk_overlap=0
                    )
                    
                    # Generate summaries and create/update retriever
                    text_summaries, table_summaries = generate_text_summaries(
                        texts, tables, summarize_texts=True
                    )
                    
                    # Handle images if present
                    img_base64_list, image_summaries = process_images(temp_dir)
                    
                    # Create or update retriever
                    if conversation_id not in retrievers:
                        vectorstore = Chroma(
                            collection_name=f"mm_rag_{conversation_id}",
                            embedding_function=OpenAIEmbeddings()
                        )
                        retrievers[conversation_id] = create_multi_vector_retriever(
                            vectorstore,
                            text_summaries,
                            texts,
                            table_summaries,
                            tables,
                            image_summaries,
                            img_base64_list
                        )
                    else:
                        # Update existing retriever with new documents
                        update_retriever(
                            retrievers[conversation_id],
                            text_summaries,
                            texts,
                            table_summaries,
                            tables,
                            image_summaries,
                            img_base64_list
                        )
                else:
                    raise HTTPException(
                        status_code=400,
                        detail=f"Unsupported file type: {file_extension}"
                    )
                    
        return {"message": "Files processed successfully"}
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/query/{conversation_id}", response_model=QueryResponse)
async def query_documents(
    query_request: QueryRequest,
    conversation: Conversation = Depends(get_conversation)
):
    """Query documents with chat history context"""
    try:
        # Get retriever for this conversation
        if query_request.conversation_id not in retrievers:
            raise HTTPException(
                status_code=404,
                detail="No documents processed for this conversation"
            )
            
        retriever = retrievers[query_request.conversation_id]
        
        # Create chat history context
        chat_history = [
            f"{msg.role}: {msg.content}" 
            for msg in conversation.messages
        ]
        
        # Create and run RAG chain with chat history
        chain = create_multimodal_rag_chain(
            retriever,
            chat_history=chat_history
        )
        response = chain.invoke(query_request.query)
        
        # Get retrieved documents
        docs = retriever.invoke(query_request.query)
        retrieved_docs = [
            str(doc) if not looks_like_base64(doc) 
            else f"data:image/jpeg;base64,{doc}"
            for doc in docs
        ]
        
        # Update conversation history
        conversation.messages.append(
            ChatMessage(
                role="user",
                content=query_request.query,
                timestamp=datetime.now()
            )
        )
        conversation.messages.append(
            ChatMessage(
                role="assistant",
                content=response,
                timestamp=datetime.now()
            )
        )
        
        return QueryResponse(
            response=response,
            retrieved_documents=retrieved_docs
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/conversations/{conversation_id}", response_model=Conversation)
async def get_conversation_history(
    conversation: Conversation = Depends(get_conversation)
):
    """Get conversation history"""
    return conversation

@app.delete("/conversations/{conversation_id}")
async def delete_conversation(
    conversation: Conversation = Depends(get_conversation)
):
    """Delete conversation and associated resources"""
    try:
        # Clean up retriever and vectorstore
        if conversation.conversation_id in retrievers:
            retriever = retrievers[conversation.conversation_id]
            retriever.vectorstore.delete_collection()
            del retrievers[conversation.conversation_id]
        
        # Remove conversation
        del conversations[conversation.conversation_id]
        return {"message": "Conversation deleted successfully"}
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# Helper function to update existing retriever
def update_retriever(
    retriever,
    text_summaries,
    texts,
    table_summaries,
    tables,
    image_summaries,
    images
):
    """Update existing retriever with new documents"""
    doc_ids = [str(uuid.uuid4()) for _ in range(len(texts) + len(tables) + len(images))]
    current_idx = 0
    
    def add_documents(summaries, contents):
        nonlocal current_idx
        if summaries:
            summary_docs = [
                Document(
                    page_content=s,
                    metadata={"doc_id": doc_ids[current_idx + i]}
                )
                for i, s in enumerate(summaries)
            ]
            retriever.vectorstore.add_documents(summary_docs)
            retriever.docstore.mset(
                list(zip(
                    doc_ids[current_idx:current_idx + len(contents)],
                    contents
                ))
            )
            current_idx += len(contents)
    
    add_documents(text_summaries, texts)
    add_documents(table_summaries, tables)
    add_documents(image_summaries, images)

def process_csv(file_path: str, temp_dir: str):
    """Process CSV file and extract elements"""
    import pandas as pd
    df = pd.read_csv(file_path)
    return df.to_dict(orient='records')  # Convert DataFrame to a list of dictionaries

def process_text(file_path: str, temp_dir: str):
    """Process text file and extract elements"""
    with open(file_path, 'r') as file:
        return file.readlines()  # Read lines from the text file

def process_doc(file_path: str, temp_dir: str):
    """Process DOC file and extract elements"""
    from docx import Document as DocxDocument
    doc = DocxDocument(file_path)
    return [para.text for para in doc.paragraphs]  # Extract text from paragraphs

def process_image(file_path: str, temp_dir: str):
    """Process image file and extract elements"""
    from PIL import Image
    img = Image.open(file_path)
    img.save(os.path.join(temp_dir, "processed_image.jpg"))  # Save processed image
    return ["Image processed and saved."]  # Placeholder for image processing result
